# intel-monitor/backend/services/report_exporter.py
"""报告导出 — markdown 转 docx / pdf。"""

import io
import logging

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _parse_markdown_to_blocks(md: str) -> list[dict]:
    """Simple markdown parser: split into blocks {type, text, level}."""
    lines = md.strip().split("\n")
    blocks = []
    buffer = []  # paragraph buffer

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                blocks.append({"type": "paragraph", "text": "\n".join(buffer)})
                buffer = []
            continue

        # Heading detection: # Title, ## Section, ### Subsection
        if stripped.startswith("# ") and len(stripped) > 2:
            if buffer:
                blocks.append({"type": "paragraph", "text": "\n".join(buffer)})
                buffer = []
            blocks.append({"type": "heading", "level": 1, "text": stripped[2:]})
        elif stripped.startswith("## ") and len(stripped) > 3:
            if buffer:
                blocks.append({"type": "paragraph", "text": "\n".join(buffer)})
                buffer = []
            blocks.append({"type": "heading", "level": 2, "text": stripped[3:]})
        elif stripped.startswith("### ") and len(stripped) > 4:
            if buffer:
                blocks.append({"type": "paragraph", "text": "\n".join(buffer)})
                buffer = []
            blocks.append({"type": "heading", "level": 3, "text": stripped[4:]})
        # Horizontal rule
        elif stripped in ("---", "***", "___", "- - -"):
            if buffer:
                blocks.append({"type": "paragraph", "text": "\n".join(buffer)})
                buffer = []
            blocks.append({"type": "hr"})
        else:
            buffer.append(stripped)

    if buffer:
        blocks.append({"type": "paragraph", "text": "\n".join(buffer)})

    return blocks


def _add_markdown_text(paragraph, text: str):
    """Add text with basic inline formatting (**bold**, *italic*) to a paragraph."""
    import re

    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for m in pattern.finditer(text):
        # Text before the match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(11)
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(11)
        pos = m.end()

    # Remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(11)


async def export_to_docx(title: str, markdown: str) -> bytes:
    """Convert markdown report to .docx bytes using python-docx."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "SimHei"
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # Blank line
    doc.add_paragraph()

    # Parse and render markdown blocks
    blocks = _parse_markdown_to_blocks(markdown)

    for block in blocks:
        if block["type"] == "heading":
            level = block["level"]
            text = block["text"]
            para = doc.add_paragraph()
            if level == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(15)
                run.font.name = "SimHei"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            elif level == 2:
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = "SimHei"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            else:
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = "SimHei"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        elif block["type"] == "hr":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("—" * 40)
            run.font.color.rgb = RGBColor(150, 150, 150)
            run.font.size = Pt(8)

        elif block["type"] == "paragraph":
            text = block["text"]
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            _add_markdown_text(para, text)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info(f"Exported docx: {title} ({len(buf.getvalue())} bytes)")
    return buf.getvalue()


async def export_to_pdf(title: str, markdown: str) -> bytes:
    """Convert markdown report to real PDF bytes.

    Renders print-friendly HTML via headless Chrome (--print-to-pdf), which
    reuses the system Chrome install (same binary used by CDP/Playwright)
    and gives correct CJK rendering. Falls back to raw HTML bytes only if
    Chrome cannot be located (caller still receives printable content).
    """
    from datetime import datetime

    html = _markdown_to_simple_html(title, markdown)
    html_with_print = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  @media print {{ @page {{ margin: 2cm; size: A4; }} }}
  body {{ font-family: "SimSun", "宋体", serif; font-size: 12pt; line-height: 1.8; max-width: 210mm; margin: 0 auto; padding: 20px; color: #1a1a1a; }}
  h1 {{ font-family: "SimHei", "黑体", sans-serif; font-size: 18pt; text-align: center; margin-bottom: 4pt; }}
  h2 {{ font-family: "SimHei", "黑体", sans-serif; font-size: 14pt; margin-top: 24pt; }}
  h3 {{ font-family: "SimHei", "黑体", sans-serif; font-size: 12pt; }}
  p {{ text-indent: 2em; margin: 6pt 0; }}
  hr {{ border: none; border-top: 1px dashed #ccc; margin: 20pt 0; }}
  .meta {{ text-align: center; color: #666; font-size: 10pt; margin-bottom: 24pt; }}
</style>
</head>
<body>
{html}
<p class="meta" style="margin-top:40pt;">报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')} · 情报监测平台</p>
</body>
</html>"""

    chrome = _find_chrome()
    if not chrome:
        logger.warning("Chrome not found for PDF export — returning HTML bytes instead")
        return html_with_print.encode("utf-8")

    import asyncio
    import os
    import shutil
    import subprocess
    import tempfile

    def _render() -> bytes:
        tmpdir = tempfile.mkdtemp(prefix="intel_report_")
        try:
            html_path = os.path.join(tmpdir, "report.html")
            pdf_path = os.path.join(tmpdir, "report.pdf")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_with_print)
            cmd = [
                chrome,
                "--headless=new",
                "--disable-gpu",
                "--no-sandbox",
                "--no-pdf-header-footer",
                f"--print-to-pdf={pdf_path}",
                "file:///" + html_path.replace("\\", "/"),
            ]
            proc = subprocess.run(cmd, capture_output=True, timeout=60)
            if proc.returncode != 0 or not os.path.exists(pdf_path):
                raise RuntimeError(
                    f"chrome print-to-pdf failed (rc={proc.returncode}): "
                    f"{proc.stderr.decode('utf-8', errors='ignore')[-300:]}"
                )
            with open(pdf_path, "rb") as f:
                data = f.read()
            if not data.startswith(b"%PDF"):
                raise RuntimeError("chrome output is not a valid PDF")
            return data
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    try:
        pdf_bytes = await asyncio.to_thread(_render)
        logger.info(f"Exported pdf: {title} ({len(pdf_bytes)} bytes)")
        return pdf_bytes
    except Exception as e:
        logger.warning(f"PDF render failed ({e}) — falling back to HTML bytes")
        return html_with_print.encode("utf-8")


def _find_chrome() -> str | None:
    """Locate a Chrome/Chromium executable, or None."""
    import os
    import shutil

    candidates = [
        os.environ.get("CHROME_PATH", ""),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    for name in ("google-chrome", "chromium", "chromium-browser", "chrome"):
        path = shutil.which(name)
        if path:
            return path
    return None


def _markdown_to_simple_html(title: str, md: str) -> str:
    """Convert basic markdown to HTML."""
    blocks = _parse_markdown_to_blocks(md)

    html_parts = [f"<h1>{title}</h1>", '<p class="meta">编制单位：情报监测平台</p>', "<hr>"]

    for block in blocks:
        if block["type"] == "heading":
            level = block["level"]
            text = block["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html_parts.append(f"<h{level + 1}>{text}</h{level + 1}>")
        elif block["type"] == "hr":
            html_parts.append("<hr>")
        elif block["type"] == "paragraph":
            text = block["text"]
            # Basic inline formatting
            import re
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
            text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
            html_parts.append(f"<p>{text}</p>")

    return "\n".join(html_parts)
